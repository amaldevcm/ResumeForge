import uuid
from operator import index
from datetime import datetime
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from Prompts.ResumeReview import get_improvement_suggestions
from Services.UserService import get_current_user
from Models.Models import Document
from Services.PineconeService import save_vector, findBestResumes
from DB import SessionLocal, Supabase
import mimetypes
import pdfplumber
import docx
import re
import os


# Function to load files
def load_file(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.pdf':
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    elif ext == '.docx':
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use .txt, .pdf, or .docx")

# Function to truncate text to a maximum number of words
def truncate_text(text, max_words=512):
    return ' '.join(text.split()[:max_words])

# Create vector embeddings and store them
def create_embeddings(resume_text, job_text):

    # Sentence Transformer model
    model = SentenceTransformer('all-MiniLM-L6-v2')

    # Create embeddings
    resume_embedding = model.encode(resume_text, convert_to_tensor=False)
    job_embedding = model.encode(job_text, convert_to_tensor=False) 

    # Section-wise encoding of the resume
    # Split resume into sections by common headings (e.g., Experience, Education, Skills)
    section_pattern = re.compile(r'(?i)(experience|education|skills|projects|summary|certifications|achievements|contact)[\s:\n]+')
    sections = section_pattern.split(resume_text)
    section_dict = {}

    # sections alternates between content and headings
    for i in range(1, len(sections), 2):
        heading = sections[i].strip().capitalize()
        content = sections[i+1].strip()
        section_dict[heading] = content

    # Encode each section separately
    section_embeddings = {}
    for heading, content in section_dict.items():
        section_embeddings[heading] = model.encode(content, convert_to_tensor=False)

    # Calculate similarity score
    similarity = cosine_similarity(
        resume_embedding.reshape(1, -1), 
        job_embedding.reshape(1, -1)
    )[0][0]

    # Store embeddings in memory (you can later use a vector database)
    vector_store = {
        'resume': resume_embedding,
        'job': job_embedding,
        'section_embeddings': section_embeddings,
        'similarity': similarity
    }
    return vector_store

# Function to grade resume against job description
def grade_resume(resume, jobDesc, title):
    # Convert binary files to text
    resume = load_file(resume)
    jobDesc = load_file(jobDesc)

    # Grade resume
    pass

def getAllResumes():
    try:
        db = SessionLocal()
        # Retrieve all resume entries for a user
        uid = get_current_user()['id']
        documents = db.query(Document).filter(Document.user_id == uid).all()
        result = []
        for document in documents:
            result.append({
                "id": str(document.id),
                "title": document.title,
                "created_date": document.created_date,
                "updated_date": document.updated_date
            })
        return result
    except Exception as e:
        raise Exception(f"Error retrieving resumes: {str(e)}")
    finally:
        db.close()

# Function to get document by ID
def getDocumentById(document_id):
    try:
        db = SessionLocal()
        resume = db.query(Document).filter(Document.id == document_id).first()
        if resume is None:
            return None
        
        url = Supabase.storage.from_('Resumes').create_signed_url(resume.resume_url, 3600)['signedURL']  # URL valid for 1 hour

        return {
            "id": str(resume.id),
            "title": resume.title,
            "resume_url": url,
            "created_date": resume.created_date,
            "updated_date": resume.updated_date
        }
    except Exception as e:
        raise Exception(f"Error retrieving resume: {str(e)}")
    finally:
        db.close()

# Function to save document
def saveDocument(resume_url, title):
    db = SessionLocal()
    
    try:
        # Save file to Supabase Storage
        file_name = str(uuid.uuid4()) + os.path.splitext(resume_url)[-1].lower()

        # load the file content        
        with open(resume_url, 'rb') as f:
            resume_content = f.read()
        mime_type, _ = mimetypes.guess_type(resume_url)

        # Upload the file to Supabase Storage
        Supabase.storage.from_('Resumes').upload(
                            file_name, 
                            resume_content,
                            {"content-type": mime_type or "application/octet-stream"}
                        )
    except Exception as e:
        raise Exception(f"Error uploading file to Supabase: {str(e)}")
    

    try:
        resume = load_file(resume_url)

        # Create embeddings using Sentence Transformer model
        model = SentenceTransformer('all-MiniLM-L6-v2')
        resume_id = str(uuid.uuid4())

        # Create embeddings
        resume_embedding = model.encode(resume, convert_to_tensor=False)
        resume_index = save_vector(resume_embedding, metadata={"resume_id": str(resume_id), "user_id": str(get_current_user()['id'])})

        new_document = Document(
            id = resume_id,
            user_id = get_current_user()['id'],
            title = title,
            resume_text = resume,
            resume_url = file_name,
            resume_vector_id = resume_index,
            created_date = datetime.now().isoformat(),
            updated_date = datetime.now().isoformat()
        )
        db.add(new_document)
        db.commit()
        return {
            "id": str(new_document.id),
            "title": new_document.title,
            "created_date": new_document.created_date,
            "updated_date": new_document.updated_date
        }
    except Exception as e:
        db.rollback()
        raise Exception(f"Error saving document: {str(e)}")
    finally:
        db.close()

# Function to get best resumes for a JD
def getBestResumes(jd_text, top_k=3):
    resume_ids = findBestResumes(jd_text, top_k)
    resumes = []
    for resume_id in resume_ids:
        resume = getDocumentById(resume_id)
        if resume:
            resumes.append(resume)
    return resumes

